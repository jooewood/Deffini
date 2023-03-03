#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
@author: zdx
"""
from glob import glob
import os
import pandas as pd
from docx import Document
from docx.shared import Inches
from tqdm import tqdm
from utils import try_copy, build_new_folder, evaluation_one_target, copytree
# from functools import partial
from eval import main as eval_main
from plot import distribution_plot, scatter_plot
from argparse import ArgumentParser

if __name__ == "__main__":
    ap = ArgumentParser()
    ap.add_argument("--exp_name", default='cf600_h100_dp0.1_False_batch320_down')
    ap.add_argument("--exps_dir", default='/y/Aurora/Fernie/EXPS')
    ap.add_argument("--report_root", default='/y/Aurora/Fernie/Report')
    ap.add_argument("--box_type", default='vio')
    ap.add_argument("--output_dir", default=None)
    ap.add_argument("--plot", default=False, action="store_true")
    args = ap.parse_args()
    
    report_root = args.report_root
    
    exp_name = args.exp_name
    
    plot_flag = args.plot
    
    box_type = args.box_type
    
    exps_dir = args.exps_dir
    
    report_dir = os.path.join(report_root, exp_name)
    if not os.path.exists(report_dir):
        os.makedirs(report_dir)
    doc_path = os.path.join(report_dir, "key_figures_tables.docx")
    fernie_exp_dir = os.path.join(exps_dir, exp_name)
    
    figure_output_dir = os.path.join(report_dir, 'Figures')
    if not os.path.exists(figure_output_dir):
        os.makedirs(figure_output_dir)
        
    doc = Document()
    sections = doc.sections
    section = sections[0]
    section.bottom_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    def WriteDataFrameToWord(doc, df):
        t = doc.add_table(df.shape[0]+1, df.shape[1])
        t.autofit = False
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]
            run = t.cell(0,j).paragraphs[0].runs[0]
            run.font.bold = True
        
        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
    
    def collect_cv_result(dataset, fernie_exp_dir, doc, method='fernie'):
        performance = pd.read_csv(os.path.join(fernie_exp_dir,
            dataset,'performances', f'{dataset}.{method}.performance'), sep='\t')
        performance_summary = pd.read_csv(os.path.join(fernie_exp_dir, 
            dataset,'performances', f'{dataset}.{method}.performance.summary'), sep='\t')
        return performance, performance_summary
    
    doc.add_heading('Figures and Tables\n')
    doc.add_paragraph('\nFigures')
    doc.add_page_break()
    
    """
    ------------------------------------------
    Fig1. Model detail
    ------------------------------------------
    """
    doc.add_paragraph('\nFigure 1. \n The Deffini model.')
    # doc.add_picture('/y/Aurora/Fernie/Report/fig1_V1.png', width=Inches(5.0))
    doc.add_picture('/y/Aurora/Fernie/Report/fig1_v2.png', width=Inches(5.0))

    """
    ------------------------------------------
    Fig2. 3-fold cross-validation of the DUD-E
    ------------------------------------------
    """
    
    dataset='DUD-E'
    doc.add_page_break()
    doc.add_paragraph('\nFigure 2. \nPerformance comparison of Smina and Deffini in clustered three-fold cross-validation of the DUD-E.')
    files = [
        f'/y/Aurora/Fernie/output/{dataset}/Smina/performances/{dataset}.Smina.performance',
        # os.path.join(exps_dir, 'Morgan', f'{dataset}/performances/{dataset}.Morgan.performance'),
        os.path.join(exps_dir, 'GanDTI', f'{dataset}/performances/{dataset}.GanDTI.performance'),
        os.path.join(exps_dir, 'CNN', f'{dataset}/performances/{dataset}.CNN.performance'),
        os.path.join(exps_dir, 'Transformer', f'{dataset}/performances/{dataset}.Transformer.performance'),
        # os.path.join(exps_dir, 'GanDTI', f'{dataset}/performances/{dataset}.GanDTI.performance'),
        os.path.join(exps_dir, exp_name, f'{dataset}/performances/{dataset}.fernie.performance'),
        ]
    names = ['Smina', 'GanDTI', 'CNN', 'Transformer', 'Deffini']
    if plot_flag:
        distribution_plot(files, names, figure_output_dir, 'fig2', box_type)
    doc.add_picture(os.path.join(figure_output_dir, 'fig2.png'), width=Inches(5.0))
    
    
    """
    ------------------------------------------
    Fig3. 3-fold cross-validation of the Kernie
    ------------------------------------------
    """
    dataset = 'Kernie'
    doc.add_page_break()
    doc.add_paragraph('\nFigure 3. \nPerformance comparison of Smina and Deffini in clustered three-fold cross-validation of the Kernie.')
    files = [
        f'/y/Aurora/Fernie/output/{dataset}/Smina/performances/{dataset}.Smina.performance',
        # os.path.join(exps_dir, 'Morgan', f'{dataset}/performances/{dataset}.Morgan.performance'),
        os.path.join(exps_dir, 'GanDTI', f'{dataset}/performances/{dataset}.GanDTI.performance'),
        os.path.join(exps_dir, 'CNN', f'{dataset}/performances/{dataset}.CNN.performance'),
        os.path.join(exps_dir, 'Transformer', f'{dataset}/performances/{dataset}.Transformer.performance'),
        # os.path.join(exps_dir, 'GanDTI', f'{dataset}/performances/{dataset}.GanDTI.performance'),
        os.path.join(exps_dir, exp_name, f'{dataset}/performances/{dataset}.fernie.performance'),
        ]
    names = ['Smina', 'GanDTI', 'CNN', 'Transformer', 'Deffini']
    
    if plot_flag:
        distribution_plot(files, names, figure_output_dir, 'fig3', box_type)
    doc.add_picture(os.path.join(figure_output_dir, 'fig3.png'), width=Inches(5.0))

    """
    ------------------------------------------
    Increase targets Kernie kinase
    ------------------------------------------
    """
    if os.path.exists(os.path.join(fernie_exp_dir, 'increase_spilt')):
        
        dataset = "Kernie"
        family = "Kinase"
        time_ = 0
        
        input_dir = os.path.join(fernie_exp_dir, f'increase_spilt/{dataset}/{family}/{time_}')
        
        l = len(glob(os.path.join(input_dir, '*')))

        dfs = []
        
        for i in range(l):
            # i = 0
            cur_input_dir = os.path.join(input_dir, str(i))
            performance_file = glob(os.path.join(cur_input_dir, 'performances/*.performance'))[0]
            performances_dir = os.path.join(cur_input_dir, 'performances')
            if not os.path.exists(performance_file):
                scores_dir = os.path.join(cur_input_dir, 'scores')
                l = len(glob(os.path.join(scores_dir, '*.score')))
                if not len(l) < 1:
                    print(f"There is no score file in {scores_dir}")
                    break
                else:
                    eval_main(input_dir=scores_dir, output_dir=performances_dir, 
                              method_name="fernie", dataset_name="MUV")
            df = pd.read_csv(performance_file, sep='\t')
            familys = []
            targets = []
            for target in df['target']:
                targets.append(target[:3])
                target = target.lower()
                if 'kinase' in target:
                    familys.append("kinase")
                elif 'nuclear' in target:
                    familys.append("nuclear")
                elif 'protease' in target:
                    familys.append("protease")
                elif 'ppic' in target:
                    familys.append("PPI")
                elif 'rnase' in target:
                    familys.append("RNase")
                elif 'gpcr' in target:
                    familys.append('GPCR')
                elif 'chaperone' in target:
                    familys.append('chaperone')
            df['target'] = targets
            df['family'] = familys
            df['no_of_targets'] = i+1
            dfs.append(df)
            
        df_m = pd.concat(dfs)
        df_m.to_excel(os.path.join(report_dir, 
            f'{dataset}_{family}_add_targets.xlsx'), index=False)
        df_m_stat = df_m.groupby(by=["no_of_targets", "family"]).mean()
        df_m_stat = df_m_stat.applymap(round, ndigits=4)
        
        family = 'kinase'
        df = df_m_stat.query('family=="%s"' % family)
        if plot_flag:
            scatter_plot(df, 'kinase', figure_output_dir, 'fig4', frac=1./10.)
            
        doc.add_page_break()
        doc.add_paragraph(f'\nFigure 4. \nIncreasing targets on {dataset} {family}')
        doc.add_picture(os.path.join(figure_output_dir, 'fig4.png'), width=Inches(5.0))




    doc.add_page_break()
    doc.add_paragraph('\nTables')
    """
    ------------------------------------------
    Clustered 3-fold cross-validation on DUD-E
    ------------------------------------------
    """
    dataset = 'DUD-E'
    
    _, fernie_cv_summary = collect_cv_result(dataset, fernie_exp_dir, doc)
    Smina_summary = pd.read_csv(
        f'/y/Aurora/Fernie/output/{dataset}/Smina/performances/{dataset}.Smina.performance.summary',
        sep='\t'
        )
    _, CNN_cv_summary = collect_cv_result(dataset, 
                                       "/y/Aurora/Fernie/EXPS/CNN", 
                                       doc, 'CNN')
    _, Transformer_cv_summary = collect_cv_result(dataset, 
                                       "/y/Aurora/Fernie/EXPS/Transformer", 
                                       doc, 'Transformer')
    # _, Morgan_cv_summary = collect_cv_result(dataset, 
    #                                    "/y/Aurora/Fernie/EXPS/Morgan", 
    #                                    doc, 'Morgan')
    _, GanDTI_cv_summary = collect_cv_result(dataset, 
                                       "/y/Aurora/Fernie/EXPS/GanDTI", 
                                       doc, 'GanDTI')
    
    df = pd.concat([Smina_summary, GanDTI_cv_summary, CNN_cv_summary, 
                    Transformer_cv_summary, fernie_cv_summary]) # GanDTI_summary,
    df = df[['AUC_ROC', 'AUC_PRC','EF1%', 'EF5%', 'EF10%']]
    df = df.T
    df.reset_index(drop=False, inplace=True)
    df.columns = ['Metric', 'Smina', 'GanDTI', 'CNN', 'Transformer', 'Deffini']
    doc.add_paragraph(f'\nCross validation on {dataset}')
    WriteDataFrameToWord(doc, df)
    
    
    """
    ------------------------------------------
    Clustered 3-fold cross-validation on Kernie
    ------------------------------------------
    """
    dataset = 'Kernie'
    
    _, fernie_cv_summary = collect_cv_result(dataset, fernie_exp_dir, doc)
    Smina_summary = pd.read_csv(
        f'/y/Aurora/Fernie/output/{dataset}/Smina/performances/{dataset}.Smina.performance.summary',
        sep='\t'
        )
    
    _, CNN_cv_summary = collect_cv_result(dataset, 
                                       "/y/Aurora/Fernie/EXPS/CNN", 
                                       doc, 'CNN')
    _, Transformer_cv_summary = collect_cv_result(dataset, 
                                       "/y/Aurora/Fernie/EXPS/Transformer", 
                                       doc, 'Transformer')
    _, Morgan_cv_summary = collect_cv_result(dataset, 
                                       "/y/Aurora/Fernie/EXPS/Morgan", 
                                       doc, 'Morgan')
    _, GanDTI_cv_summary = collect_cv_result(dataset, 
                                       "/y/Aurora/Fernie/EXPS/GanDTI", 
                                       doc, 'GanDTI')
    
    df = pd.concat([Smina_summary, GanDTI_cv_summary, CNN_cv_summary, 
                    Transformer_cv_summary, fernie_cv_summary]) 
    df = df[['AUC_ROC', 'AUC_PRC','EF1%', 'EF5%', 'EF10%']]
    df = df.T
    df.reset_index(drop=False, inplace=True)
    df.columns = ['Metric', 'Smina', 'GanDTI', 'CNN', 'Transformer', 'Deffini']
    doc.add_page_break()
    doc.add_paragraph(f'\nCross validation on {dataset}')
    WriteDataFrameToWord(doc, df)

    """
    ------------------------------------------
    Testing on MUV
    ------------------------------------------
    """
    def collect_MUV_scores(method='fernie'):
        if method=='fernie':
            exp_dir = fernie_exp_dir
            method = 'Deffini'
        else:
            exp_dir = os.path.join('/y/Aurora/Fernie/EXPS/', method)
        MUV_dirs = os.listdir(os.path.join(exp_dir, 'MUV'))
        output_dir = os.path.join(exp_dir, 'MUV_scores')
        build_new_folder(output_dir)
        for training_set in tqdm(MUV_dirs):
            # training_set = MUV_dirs[1]
            dir_path = os.path.join(exp_dir, 'MUV', training_set)
            current_dir = os.path.join(dir_path, 'scores')
            src_scores_names = os.listdir(current_dir)
            for file in src_scores_names:
                # file = src_scores_names[0]
                if method == 'Deffini':
                    target_name, target_family = file.split('.')[0:2]
                else:
                    target_family, target_name = file.split('.')[0:2]
                new_name = '.'.join([target_name, target_family, method, 
                                     training_set.replace('.', '_'), 'score'])
                try_copy(os.path.join(current_dir, file), os.path.join(output_dir, new_name))
    # collect_MUV_scores()
    # collect_MUV_scores('Morgan')
    # collect_MUV_scores('CNN')
    # collect_MUV_scores('Transformer')
    # collect_MUV_scores('GanDTI')
    
    fernie_MUV_scores = glob(os.path.join(os.path.join(fernie_exp_dir, 'MUV_scores'), 
                                          '*.score'))
    # Morgan_MUV_scores = glob(os.path.join('/y/Aurora/Fernie/EXPS/Morgan/MUV_scores', '*.score'))
    GanDTI_MUV_scores = glob(os.path.join('/y/Aurora/Fernie/EXPS/GanDTI/MUV_scores', '*.score'))
    CNN_MUV_scores = glob(os.path.join('/y/Aurora/Fernie/EXPS/CNN/MUV_scores', '*.score'))
    Transformer_MUV_scores = glob(os.path.join('/y/Aurora/Fernie/EXPS/Transformer/MUV_scores', '*.score'))
    Smina_MUV_scores = glob(os.path.join('/y/Aurora/Fernie/output/MUV/scores/Smina', '*.score'))
    
    # lf_transformer_scores = glob(os.path.join('/y/Aurora/Fernie/output/MUV/scores/transformer', '*.score'))
    
    all_scores = fernie_MUV_scores + Smina_MUV_scores + GanDTI_MUV_scores +\
        CNN_MUV_scores + Transformer_MUV_scores
    
    target_names = []
    target_familys = []
    methods = []
    training_sets = []
    auc_rocs = []
    auc_prcs = []
    ef1s = []
    ef5s = []
    ef10s = []
    for file in tqdm(all_scores):
        # file = all_scores[0]
        auc_roc, auc_prc, ef1, ef5, ef10 = evaluation_one_target(file)
        target_name, target_family, method, training_set, _ = os.path.basename(file).split('.')
        
        target_names.append(target_name)
        target_familys.append(target_family)
        methods.append(method)
        training_sets.append(training_set)
        auc_rocs.append(auc_roc)
        auc_prcs.append(auc_prc)
        ef1s.append(ef1)
        ef5s.append(ef5)
        ef10s.append(ef10)
        
    df = pd.DataFrame({
        'target' : target_names,
        'family' : target_familys,
        'method' : methods,
        'training_set' : training_sets,
        'AUC_ROC' : auc_rocs,
        'AUC_PRC' : auc_prcs,
        'EF1%' : ef1s,
        'EF5%' : ef5s,
        'EF10%' : ef10s
        })
    
    # x = []
    # for y in df.family:
    #     if y == 'nuclear':
    #         y = 'Nuclear'
    #     if y == 'PPIc':
    #         y = 'PPI'
    #     if y == 'protease':
    #         y = 'Protease'
    #     if y == 'kinase':
    #         y = 'Kinase'
    #     if y == 'Rnase':
    #         y = 'RNase'
    #     x.append(y)
    # df['family'] = x

    df.drop_duplicates(inplace=True)
    tmp = []
    for x in df['family']:
        x = x.lower()
        if x == "ppic" or x == 'ppi':
            x = "PPI"
        if x == 'gpcr':
            x = "GPCR"
        if x == "rnase":
            x = "RNase"
        tmp.append(x)
    df['family'] = tmp
    
    
    df.to_excel(os.path.join(report_dir, 'MUV_test_details.xlsx'), index=False)
    
    df_stat_f_m_t = df.groupby(by=["family", "method", "training_set"]).mean()
    df_stat_f_m_t = df_stat_f_m_t.applymap(round, ndigits=4)
    counts = df.groupby(["family", "method", "training_set"]).size()
    df_stat_f_m_t['count'] = counts
    
    """
    family  method
    """
    
    familys = list(set(df.family))
    methods = list(set(df.method))
    training_sets = list(set(df.training_set))
    dfs = []
    for family in familys:
        for method in methods:
            # for training_set in training_sets:
            sub = df_stat_f_m_t.query('family == "%s" & method == "%s" ' % (family, method))
            sub.sort_values(['AUC_ROC', 'AUC_PRC', 'EF1%', 'EF5%', 'EF10%'], ascending=False, inplace=True)
            dfs.append(sub)
    
    df_stat_f_m_t_sort = pd.concat(dfs)
    df_stat_f_m_t_sort.to_excel(os.path.join(report_dir, 'MUV_test_family_statistic.xlsx'))
    
    kinase_Deffini = df_stat_f_m_t_sort.query('family=="kinase"')
    kinase_Deffini = kinase_Deffini.query('method=="Smina" | method=="Deffini"')
    kinase_Deffini = kinase_Deffini.query('training_set in ["DUD-E_Protease", "DUD-E", "Kernie-MUV", "DUD-E_Miscellaneous", "Kernie", "DUD-E_Kinase", "DUD-E_GPCR"]')
    kinase_Deffini.reset_index(drop=False, inplace=True)
    kinase_Deffini.drop(columns=['family', 'method', 'count'], inplace=True)
    kinase_Deffini.sort_values(['AUC_ROC'], ascending=False, inplace=True)
    doc.add_page_break()
    doc.add_paragraph('\n Comparing the performance of Deffini models trained with different subsets of DUD-E on the MUV_kinase subset.')
    WriteDataFrameToWord(doc, kinase_Deffini)

    protease_Deffini = df_stat_f_m_t_sort.query('family=="protease"')
    protease_Deffini = protease_Deffini.query('method=="Smina" | method=="Deffini"')
    protease_Deffini = protease_Deffini.query('training_set in ["DUD-E_Protease", "DUD-E", "Kernie-MUV", "DUD-E_Miscellaneous", "Kernie", "DUD-E_Kinase", "DUD-E_GPCR"]')
    protease_Deffini.reset_index(drop=False, inplace=True)
    protease_Deffini.sort_values(['AUC_ROC'], ascending=False, inplace=True)
    protease_Deffini.drop(columns=['family', 'method', 'count'], inplace=True)
    doc.add_page_break()
    doc.add_paragraph('\n Comparing the performance of Deffini models trained with different subsets of DUD-E on the MUV_protease subset.')
    WriteDataFrameToWord(doc, protease_Deffini)


    
    """
    method  training_set
    """
    df_stat_m_t = df.groupby(by=["method", "training_set"]).mean()
    df_stat_m_t = df_stat_m_t.applymap(round, ndigits=4)
    counts = df.groupby(["method", "training_set"]).size()
    df_stat_m_t['count'] = counts
    
    methods = list(set(df.method))
    training_sets = list(set(df.training_set))
    dfs = []
    for method in methods:
        # for training_set in training_sets:
        sub = df_stat_m_t.query('method == "%s" ' % (method))
        sub.sort_values(['AUC_ROC', 'AUC_PRC', 'EF1%', 'EF5%', 'EF10%'], ascending=False, inplace=True)
        dfs.append(sub)
    df_stat_m_t = pd.concat(dfs)
    df_stat_m_t.to_excel(os.path.join(report_dir, 'MUV_test_method_statistic.xlsx'))
    
    DUDE_MUV = df_stat_m_t.query('training_set == "DUD-E" | training_set == "None" ')
    DUDE_MUV.reset_index(drop=False, inplace=True)
    DUDE_MUV.drop(columns=['training_set', 'count'], inplace=True)
    DUDE_MUV.sort_values(['EF1%'], ascending=False, inplace=True)
    # DUDE_MUV.to_excel(os.path.join(report_dir, 'DUD-E_MUV_test_method.xlsx'), index=False)
    doc.add_page_break()
    doc.add_paragraph('\n Comparing the performance of different models trained with entire DUD-E and Smina on the whole MUV dataset.')
    WriteDataFrameToWord(doc, DUDE_MUV)

    
    Kernie_MUV = df_stat_m_t.query('training_set == "Kernie" | training_set == "None" ')
    Kernie_MUV.reset_index(drop=False, inplace=True)
    Kernie_MUV.drop(columns=['training_set', 'count'], inplace=True)
    Kernie_MUV.sort_values(['EF1%'], ascending=False, inplace=True)
    # Kernie_MUV.to_excel(os.path.join(report_dir, 'Kernie_MUV_test_method.xlsx'), index=False)
    doc.add_page_break()
    doc.add_paragraph('\n Comparing the performance of different models trained with entire Kernie and Smina on the whole MUV dataset.')
    WriteDataFrameToWord(doc, Kernie_MUV)
    doc.save(doc_path)

    """
    target 
    """
    targets = list(set(df.target))
    dfs = []
    for target in targets:
        # for training_set in training_sets:
        sub = df.query('target == "%s" ' % (target))
        sub.sort_values(['AUC_ROC', 'AUC_PRC', 'EF1%', 'EF5%', 'EF10%'], ascending=False, inplace=True)
        dfs.append(sub)
    df_stat_tar = pd.concat(dfs)
    df_stat_tar.set_index(['target', 'family', 'method'], inplace=True)
    df_stat_tar.to_excel(os.path.join(report_dir, 'MUV_test_target_statistic.xlsx'))


    if args.output_dir is not None:
        gitlab_ci_dir = os.path.join(args.output_dir, exp_name)
        if not os.path.exists(gitlab_ci_dir):
            os.makedirs(gitlab_ci_dir)
        copytree(report_dir, gitlab_ci_dir)

